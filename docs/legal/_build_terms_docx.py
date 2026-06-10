# -*- coding: utf-8 -*-
"""Generate a polished, lawyer-ready .docx from syarat-ketentuan.md."""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "syarat-ketentuan.md")
OUT = os.path.join(HERE, "syarat-ketentuan.docx")

INK = RGBColor(0x1F, 0x2A, 0x44)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---------- Page setup ----------
sec = doc.sections[0]
sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.4)
sec.left_margin = Cm(2.6); sec.right_margin = Cm(2.6)
sec.different_first_page_header_footer = True

# ---------- Base styles ----------
st = doc.styles
normal = st["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
pf = normal.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.line_spacing = 1.18; pf.space_after = Pt(6)

h1 = st["Heading 1"]
h1.font.name = "Calibri"; h1.font.size = Pt(12.5); h1.font.bold = True; h1.font.color.rgb = INK
h1.paragraph_format.space_before = Pt(16); h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.keep_with_next = True

h2 = st["Heading 2"]
h2.font.name = "Calibri"; h2.font.size = Pt(11); h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
h2.paragraph_format.space_before = Pt(10); h2.paragraph_format.space_after = Pt(3)
h2.paragraph_format.keep_with_next = True


def add_runs(p, text):
    for i, seg in enumerate(text.split("**")):
        if not seg:
            continue
        r = p.add_run(seg)
        if i % 2 == 1:
            r.bold = True


def list_item(prefix, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.85); fmt.first_line_indent = Cm(-0.85)
    fmt.space_after = Pt(3); fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.tab_stops.add_tab_stop(Cm(0.85))
    p.add_run(prefix + "\t")
    add_runs(p, text)
    return p


def field(paragraph, instr, placeholder=""):
    r = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); r._r.append(b)
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr; r._r.append(t)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate'); r._r.append(s)
    if placeholder:
        ph = OxmlElement('w:t'); ph.text = placeholder; r._r.append(ph)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); r._r.append(e)


# ===================== COVER =====================
doc.add_paragraph().paragraph_format.space_before = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FIN.WEDDING"); r.font.size = Pt(15); r.bold = True; r.font.color.rgb = INK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Platform Undangan Pernikahan Digital"); r.font.size = Pt(10); r.font.color.rgb = GREY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run("SYARAT DAN KETENTUAN"); r.font.size = Pt(26); r.bold = True; r.font.color.rgb = INK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Terms of Service"); r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_before = Pt(18)
tbl = doc.add_table(rows=0, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light List Accent 1"
meta = [
    ("Nama Dokumen", "Syarat & Ketentuan Layanan"),
    ("Penyelenggara", "FIN.WEDDING"),
    ("Versi", "1.0 (Draf)"),
    ("Tanggal Berlaku", "[TANGGAL BERLAKU]"),
    ("Yurisdiksi", "Republik Indonesia"),
    ("Status", "DRAF - Menunggu Peninjauan Penasihat Hukum"),
]
for k, v in meta:
    cells = tbl.add_row().cells
    rr = cells[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(10)
    cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
for row in tbl.rows:
    row.cells[0].width = Cm(4.5); row.cells[1].width = Cm(9.5)

note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_before = Pt(28)
r = note.add_run("Dokumen ini merupakan perjanjian yang mengikat secara hukum. "
                 "Harap dibaca dengan saksama.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
doc.add_page_break()

# ===================== TOC =====================
p = doc.add_paragraph(); r = p.add_run("DAFTAR ISI")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = INK
p.paragraph_format.space_after = Pt(10)
toc_p = doc.add_paragraph()
field(toc_p, 'TOC \\o "1-2" \\h \\z \\u',
      "Buka di Microsoft Word, klik kanan di sini lalu pilih Update Field untuk memuat Daftar Isi.")
doc.add_page_break()

# ===================== BODY =====================
lines = open(SRC, encoding="utf-8").read().split("\n")
num_re = re.compile(r"^(\d+)\.\s+(.*)")

for raw in lines:
    s = raw.strip()
    if s == "" or s == "---" or s.startswith("# "):
        continue
    if s.startswith("_") and s.endswith("_") and len(s) > 2:
        continue
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=2)
    elif s.startswith("## "):
        body = s[3:].strip()
        m = num_re.match(body)
        doc.add_heading(f"Pasal {m.group(1)} - {m.group(2)}" if m else body, level=1)
    elif s.startswith("- "):
        list_item("•", s[2:].strip())
    elif num_re.match(s):
        m = num_re.match(s)
        list_item(f"{m.group(1)}.", m.group(2).strip())
    else:
        add_runs(doc.add_paragraph(), s)

# ===================== ANNEX =====================
doc.add_page_break()
doc.add_heading("Lampiran A - Catatan untuk Peninjau Hukum", level=1)
p = doc.add_paragraph()
r = p.add_run("Bagian ini BUKAN bagian dari Syarat & Ketentuan yang dipublikasikan kepada Pengguna. "
              "Lampiran ini disiapkan untuk membantu asisten/penasihat hukum meninjau dokumen sebelum "
              "go-live. Hapus seluruh Lampiran ini pada versi final yang dipublikasikan.")
r.italic = True; r.font.color.rgb = GREY

doc.add_heading("A.1 Placeholder yang wajib dilengkapi", level=2)
pt = doc.add_table(rows=1, cols=2); pt.style = "Light Grid Accent 1"
for c, txt in zip(pt.rows[0].cells, ("Placeholder", "Diisi dengan")):
    rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.size = Pt(9.5)
for a, b in [
    ("[ALAMAT]", "Alamat domisili hukum badan usaha penyelenggara"),
    ("[EMAIL]", "Alamat email resmi untuk pemberitahuan & aduan"),
    ("[NOMOR WHATSAPP]", "Nomor kontak resmi layanan pelanggan"),
    ("[DOMAIN]", "Domain final layanan (mis. finwedding.id)"),
    ("[TANGGAL BERLAKU]", "Tanggal dokumen mulai berlaku"),
    ("[PENGADILAN NEGERI]", "Pengadilan Negeri / forum (atau BANI) sesuai domisili"),
]:
    cells = pt.add_row().cells
    cells[0].paragraphs[0].add_run(a).font.size = Pt(9.5)
    cells[1].paragraphs[0].add_run(b).font.size = Pt(9.5)

doc.add_heading("A.2 Klausul yang memerlukan peninjauan hukum (prioritas)", level=2)
for t, d in [
    ("Pasal 9.3 & 9.4 - Pembatasan Tanggung Jawab & Indemnifikasi",
     "Berpotensi berbenturan dengan Pasal 18 UU No. 8/1999 tentang Perlindungan Konsumen (larangan "
     "klausula baku yang mengalihkan tanggung jawab pelaku usaha). Klausul yang terlalu luas dapat "
     "dinyatakan batal demi hukum - perlu dipersempit."),
    ("Pasal 10.3 - Penghentian Layanan & refund prorata",
     "Komitmen pengembalian dana prorata bila layanan dihentikan permanen; pastikan selaras kemampuan "
     "finansial dan Kebijakan Pengembalian Dana."),
    ("Pasal 11.2 - Penyelesaian Sengketa",
     "Tetapkan forum final (Pengadilan Negeri domisili atau arbitrase BANI) sesuai badan usaha."),
    ("Pasal 8.5 - Pengembalian Dana",
     "Pastikan kebijakan non-refundable tidak menutup hak konsumen atas produk cacat/tidak sesuai (UUPK)."),
    ("Status badan usaha & PSE",
     "Tentukan bentuk badan usaha (PT/CV/perorangan) dan tinjau kewajiban pendaftaran PSE Lingkup Privat "
     "(Permenkominfo 5/2020) serta aspek perpajakan."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.85); p.paragraph_format.first_line_indent = Cm(-0.85)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.85))
    rr = p.add_run("•\t" + t + ". "); rr.bold = True
    p.add_run(d)

doc.add_heading("A.3 Konsistensi dokumen", level=2)
doc.add_paragraph("Samakan tanggal berlaku dan identitas penyelenggara pada ketiga dokumen legal: "
                  "Syarat & Ketentuan, Kebijakan Privasi, dan Kebijakan Pengembalian Dana.")

# ===================== HEADER / FOOTER =====================
hp = sec.header.paragraphs[0]
hr = hp.add_run("FIN.WEDDING  ·  Syarat & Ketentuan")
hr.font.size = Pt(8); hr.font.color.rgb = GREY

fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("Halaman ")
field(fp, "PAGE", "1")
fp.add_run(" dari ")
field(fp, "NUMPAGES", "1")
for rr in fp.runs:
    rr.font.size = Pt(8); rr.font.color.rgb = GREY

upd = OxmlElement('w:updateFields'); upd.set(qn('w:val'), 'true')
doc.settings.element.append(upd)

doc.save(OUT)
print("SAVED:", OUT)
print("size KB:", round(os.path.getsize(OUT) / 1024, 1))
